#!/usr/bin/env python3
"""Generate Word document: Swiss Parliament OData API Specification."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_field_table(doc, fields, caption=None):
    """Add a formatted table with field definitions."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, text in enumerate(['Attribut', 'Typ', 'Pflicht', 'Beschreibung']):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for name, typ, req, desc in fields:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = typ
        row[2].text = req
        row[3].text = desc
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacing

def add_endpoint_box(doc, method, path, description):
    p = doc.add_paragraph()
    run = p.add_run(f'{method}  ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    run.font.size = Pt(10)
    run = p.add_run(path)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p2 = doc.add_paragraph(description)
    p2.paragraph_format.space_after = Pt(2)

def add_query_params_table(doc, params):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light List Accent 1'
    hdr = table.rows[0].cells
    for i, text in enumerate(['Parameter', 'Beispiel', 'Beschreibung']):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for name, example, desc in params:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = example
        row[2].text = desc
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Swiss Parliament OData API', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Vollständige API-Spezifikation', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n')
p.add_run('Parlamentarisches Informationssystem (ParIS)').bold = True
p.add_run('\n')
p.add_run('https://ws.parlament.ch/odata.svc').font.name = 'Consolas'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run(f'\nErstellt am: {datetime.date.today().strftime("%d.%m.%Y")}')

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Inhaltsverzeichnis', level=1)
toc_items = [
    '1. Übersicht',
    '2. Allgemeine Konventionen',
    '3. Datenformate',
    '4. Paginierung',
    '5. Endpoints',
    '   5.1 Session (Sessionen)',
    '   5.2 Business (Geschäfte)',
    '   5.3 BusinessRoles (Urheber)',
    '   5.4 MemberCouncil (Parlamentarier)',
    '   5.5 Meeting (Sitzungen / Tagesordnung)',
    '   5.6 Subject (Tagesordnungspunkte)',
    '   5.7 SubjectBusiness (Verknüpfung Tagesordnung–Geschäft)',
    '   5.8 Transcript (Wortmeldungen)',
    '   5.9 Vote (Abstimmungen)',
    '   5.10 Voting (Stimmabgaben)',
    '   5.11 PersonInterest (Interessensbindungen)',
    '   5.12 PersonOccupation (Berufe / Tätigkeiten)',
    '6. Datenmodelle (DTOs)',
    '7. OData Response-Wrapper',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    if not item.startswith('   '):
        for r in p.runs:
            r.bold = True

doc.add_page_break()

# ============================================================
# 1. ÜBERSICHT
# ============================================================
doc.add_heading('1. Übersicht', level=1)
doc.add_paragraph(
    'Die Swiss Parliament OData API (ParIS) stellt legislative Daten des Schweizer Parlaments '
    'als OData v4 Service bereit. Die API ermöglicht den Zugriff auf Sessionen, parlamentarische '
    'Geschäfte, Parlamentarier, Abstimmungen, Wortmeldungen und weitere Daten.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Base URL: ').bold = True
run = p.add_run('https://ws.parlament.ch/odata.svc')
run.font.name = 'Consolas'

p = doc.add_paragraph()
p.add_run('Protokoll: ').bold = True
p.add_run('OData v4 mit JSON-Format')

p = doc.add_paragraph()
p.add_run('Authentifizierung: ').bold = True
p.add_run('Keine (öffentlich zugänglich)')

p = doc.add_paragraph()
p.add_run('HTTP-Methode: ').bold = True
p.add_run('GET (nur lesender Zugriff)')

# ============================================================
# 2. ALLGEMEINE KONVENTIONEN
# ============================================================
doc.add_heading('2. Allgemeine Konventionen', level=1)

doc.add_heading('Sprachfilter', level=3)
doc.add_paragraph(
    'Alle Abfragen verwenden den Sprachfilter Language eq \'DE\' für deutschsprachige Resultate. '
    'Die Sprache wird je nach Endpoint entweder im Composite Key (z.B. Session(ID=5115,Language=\'DE\')) '
    'oder im $filter-Parameter angegeben.'
)

doc.add_heading('Composite Keys', level=3)
doc.add_paragraph(
    'Viele Endpoints verwenden zusammengesetzte Schlüssel aus ID und Sprache:'
)
items = [
    '/Session(ID=<id>,Language=\'DE\')',
    '/Business(ID=<id>,Language=\'DE\')',
    '/MemberCouncil(ID=<id>,Language=\'DE\')',
    '/Meeting(ID=<id>L,Language=\'DE\')  — L-Suffix für Long-Typ',
    '/Subject(ID=<id>L,Language=\'DE\')  — L-Suffix für Long-Typ',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Consolas'
        r.font.size = Pt(9)

doc.add_heading('OData Query-Optionen', level=3)
add_field_table(doc, [
    ('$filter', 'String', 'Ja', 'Filterbedingungen (eq, gt, and, or)'),
    ('$select', 'String', 'Nein', 'Kommagetrennte Liste der zurückzugebenden Felder'),
    ('$orderby', 'String', 'Nein', 'Sortierung (Feldname asc/desc)'),
    ('$top', 'Int', 'Nein', 'Maximale Anzahl Ergebnisse'),
    ('$format', 'String', 'Ja', 'Antwortformat, immer "json"'),
])

# ============================================================
# 3. DATENFORMATE
# ============================================================
doc.add_heading('3. Datenformate', level=1)

doc.add_heading('OData DateTime', level=3)
doc.add_paragraph(
    'Datumsfelder werden im OData-Format /Date(milliseconds)/ oder /Date(milliseconds+0000)/ '
    'zurückgegeben. Beispiel: /Date(1609459200000)/ entspricht 2021-01-01T00:00:00Z. '
    'Die Millisekunden seit Unix-Epoch werden extrahiert und in ein Date-Objekt konvertiert.'
)

doc.add_heading('DateTime in Filtern', level=3)
doc.add_paragraph(
    'In $filter-Ausdrücken wird das Format datetime\'yyyy-MM-ddTHH:mm:ss\' verwendet. '
    'Beispiel: Modified gt datetime\'2024-01-15T08:30:00\''
)

doc.add_heading('Long-Typ (L-Suffix)', level=3)
doc.add_paragraph(
    'Einige IDs (SubjectID, MeetingID) erfordern ein L-Suffix im Filter, '
    'um den OData Long-Typ (Int64) zu kennzeichnen. Beispiel: IdSubject eq 123456L'
)

# ============================================================
# 4. PAGINIERUNG
# ============================================================
doc.add_heading('4. Paginierung', level=1)
doc.add_paragraph(
    'Die API verwendet automatische Paginierung über das __next-Feld in der OData-Response. '
    'Wenn mehr Ergebnisse vorhanden sind als auf einer Seite zurückgegeben werden, enthält '
    'die Response einen __next-Link zur nächsten Seite. Die Seitengrösse wird serverseitig '
    'bestimmt (typisch 100–1000 Einträge). Die Implementation folgt der __next-Kette '
    'bis alle Resultate gesammelt sind (fetchAllPages).'
)

# ============================================================
# 5. ENDPOINTS
# ============================================================
doc.add_heading('5. Endpoints', level=1)

# --- 5.1 Session ---
doc.add_heading('5.1 Session (Sessionen)', level=2)

doc.add_heading('Alle Sessionen abrufen', level=3)
add_endpoint_box(doc, 'GET', '/Session', 'Ruft die Liste aller parlamentarischen Sessionen ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE'", 'Sprachfilter (obligatorisch)'),
    ('$orderby', 'StartDate desc', 'Sortierung nach Startdatum absteigend'),
    ('$top', '20', 'Begrenzung auf die letzten 20 Sessionen'),
    ('$format', 'json', 'Antwortformat'),
])

doc.add_heading('Einzelne Session abrufen', level=3)
add_endpoint_box(doc, 'GET', "/Session(ID=<sessionID>,Language='DE')",
                 'Ruft eine einzelne Session anhand ihrer ID ab.')

doc.add_heading('Navigations-Properties', level=3)
doc.add_paragraph('Von einer Session aus können verknüpfte Geschäfte abgerufen werden:')
p = doc.add_paragraph()
run = p.add_run("/Session(ID=<sessionID>,Language='DE')/Businesses")
run.font.name = 'Consolas'
run.font.size = Pt(10)

# --- 5.2 Business ---
doc.add_heading('5.2 Business (Geschäfte)', level=2)

doc.add_heading('Geschäfte einer Session', level=3)
add_endpoint_box(doc, 'GET', "/Session(ID=<sessionID>,Language='DE')/Businesses",
                 'Ruft alle Geschäfte einer bestimmten Session ab.')
add_query_params_table(doc, [
    ('$format', 'json', 'Antwortformat'),
    ('$select', 'ID,Title,BusinessShortNumber,...', 'Feldauswahl (siehe DTO)'),
])

doc.add_heading('Geschäfte mit inkrementeller Synchronisation', level=3)
add_endpoint_box(doc, 'GET', "/Session(ID=<sessionID>,Language='DE')/Businesses",
                 'Ruft nur seit einem bestimmten Datum geänderte Geschäfte ab.')
add_query_params_table(doc, [
    ('$filter', "Modified gt datetime'2024-01-15T08:30:00'", 'Nur Änderungen seit Datum'),
    ('$select', 'ID,Title,...', 'Feldauswahl'),
    ('$format', 'json', 'Antwortformat'),
])

doc.add_heading('Einzelnes Geschäft', level=3)
add_endpoint_box(doc, 'GET', "/Business(ID=<businessID>,Language='DE')",
                 'Ruft ein einzelnes Geschäft anhand seiner ID ab.')
add_query_params_table(doc, [
    ('$format', 'json', 'Antwortformat'),
    ('$select', 'ID,Title,BusinessShortNumber,...', 'Feldauswahl'),
])

# --- 5.3 BusinessRoles ---
doc.add_heading('5.3 BusinessRoles (Urheber)', level=2)
add_endpoint_box(doc, 'GET', "/Business(ID=<businessID>,Language='DE')/BusinessRoles",
                 'Ruft die Rollen (z.B. Urheber) eines Geschäfts ab.')
add_query_params_table(doc, [
    ('$filter', 'Role eq 7', 'Rolle 7 = Urheber/Originator'),
    ('$select', 'ID,Role,RoleName,MemberCouncilNumber,BusinessNumber', 'Feldauswahl'),
    ('$format', 'json', 'Antwortformat'),
])

# --- 5.4 MemberCouncil ---
doc.add_heading('5.4 MemberCouncil (Parlamentarier)', level=2)

doc.add_heading('Alle aktiven Parlamentarier', level=3)
add_endpoint_box(doc, 'GET', '/MemberCouncil',
                 'Ruft alle aktiven Mitglieder des National- und Ständerats ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and Active eq true", 'Nur aktive, deutschsprachig'),
    ('$select', 'ID,PersonNumber,FirstName,LastName,...', 'Feldauswahl'),
    ('$orderby', 'LastName,FirstName', 'Alphabetische Sortierung'),
    ('$format', 'json', 'Antwortformat'),
])

doc.add_heading('Einzelner Parlamentarier (Basis)', level=3)
add_endpoint_box(doc, 'GET', "/MemberCouncil(ID=<personNumber>,Language='DE')",
                 'Ruft Basisdaten eines einzelnen Parlamentariers ab.')

doc.add_heading('Einzelner Parlamentarier (Detail)', level=3)
add_endpoint_box(doc, 'GET', "/MemberCouncil(ID=<personNumber>,Language='DE')",
                 'Ruft erweiterte Daten eines Parlamentariers ab (zusätzliche Felder via $select).')
doc.add_paragraph(
    'Durch Erweiterung der $select-Liste können Detailfelder wie Geburtsdatum, Zivilstand, '
    'Militärrang, Beitritts-/Austrittsdatum etc. abgerufen werden.'
)

# --- 5.5 Meeting ---
doc.add_heading('5.5 Meeting (Sitzungen / Tagesordnung)', level=2)
add_endpoint_box(doc, 'GET', '/Meeting',
                 'Ruft die Sitzungen (Tagesordnung) einer Session ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and IdSession eq <sessionID>", 'Session-Filter'),
    ('$select', 'ID,MeetingNumber,Date,Begin,...', 'Feldauswahl'),
    ('$orderby', 'Date,SortOrder', 'Chronologische Sortierung'),
    ('$format', 'json', 'Antwortformat'),
])

# --- 5.6 Subject ---
doc.add_heading('5.6 Subject (Tagesordnungspunkte)', level=2)
add_endpoint_box(doc, 'GET', "/Meeting(ID=<meetingID>L,Language='DE')/Subjects",
                 'Ruft die Tagesordnungspunkte einer Sitzung ab.')
add_query_params_table(doc, [
    ('$select', 'ID,IdMeeting,SortOrder,VerbalixOid', 'Feldauswahl'),
    ('$orderby', 'SortOrder', 'Sortierung nach Reihenfolge'),
    ('$format', 'json', 'Antwortformat'),
])

# --- 5.7 SubjectBusiness ---
doc.add_heading('5.7 SubjectBusiness (Verknüpfung Tagesordnung–Geschäft)', level=2)

doc.add_heading('SubjectBusiness nach Geschäftsnummer', level=3)
add_endpoint_box(doc, 'GET', '/SubjectBusiness',
                 'Findet die Subject-IDs, die mit einem bestimmten Geschäft verknüpft sind.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and BusinessNumber eq <businessID>", 'Geschäftsfilter'),
    ('$select', 'IdSubject', 'Nur Subject-ID'),
    ('$format', 'json', 'Antwortformat'),
])

doc.add_heading('SubjectBusiness nach Subject', level=3)
add_endpoint_box(doc, 'GET', "/Subject(ID=<subjectID>L,Language='DE')/SubjectsBusiness",
                 'Ruft die verknüpften Geschäfte eines Tagesordnungspunkts ab.')
add_query_params_table(doc, [
    ('$select', 'IdSubject,BusinessNumber,BusinessShortNumber,Title,SortOrder', 'Feldauswahl'),
    ('$format', 'json', 'Antwortformat'),
])

# --- 5.8 Transcript ---
doc.add_heading('5.8 Transcript (Wortmeldungen)', level=2)
add_endpoint_box(doc, 'GET', '/Transcript',
                 'Ruft die Wortmeldungen (Redebeiträge) zu einem Tagesordnungspunkt ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and IdSubject eq <idSubject>L", 'Subject-Filter (L-Suffix!)'),
    ('$select', 'ID,PersonNumber,SpeakerFullName,Text,...', 'Feldauswahl'),
    ('$orderby', 'MeetingDate,SortOrder', 'Chronologische Sortierung'),
    ('$format', 'json', 'Antwortformat'),
])

# --- 5.9 Vote ---
doc.add_heading('5.9 Vote (Abstimmungen)', level=2)

doc.add_heading('Alle Abstimmungen einer Session', level=3)
add_endpoint_box(doc, 'GET', '/Vote',
                 'Ruft alle Abstimmungen einer Session ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and IdSession eq <sessionID>", 'Session-Filter'),
    ('$select', 'ID,BusinessNumber,BillTitle,MeaningYes,MeaningNo,...', 'Feldauswahl'),
    ('$format', 'json', 'Antwortformat'),
])

doc.add_heading('Inkrementelle Synchronisation', level=3)
add_endpoint_box(doc, 'GET', '/Vote',
                 'Ruft nur geänderte Abstimmungen seit einem Datum ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and IdSession eq <id> and Modified gt datetime'...'", 'Kombinierter Filter'),
    ('$select', '...', 'Feldauswahl'),
    ('$format', 'json', 'Antwortformat'),
])

# --- 5.10 Voting ---
doc.add_heading('5.10 Voting (Stimmabgaben)', level=2)
add_endpoint_box(doc, 'GET', '/Voting',
                 'Ruft die einzelnen Stimmabgaben einer Abstimmung ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and IdVote eq <voteID>", 'Abstimmungs-Filter'),
    ('$select', 'ID,IdVote,PersonNumber,Decision,DecisionText', 'Feldauswahl'),
    ('$format', 'json', 'Antwortformat'),
])
doc.add_paragraph('Decision-Werte:')
items = [
    ('1', 'Ja'),
    ('2', 'Nein'),
    ('3', 'Enthaltung'),
    ('4', 'Nicht teilgenommen'),
    ('5', 'Entschuldigt'),
    ('6', 'Präsident'),
]
for code, meaning in items:
    doc.add_paragraph(f'{code} = {meaning}', style='List Bullet')

# --- 5.11 PersonInterest ---
doc.add_heading('5.11 PersonInterest (Interessensbindungen)', level=2)
add_endpoint_box(doc, 'GET', '/PersonInterest',
                 'Ruft die Interessensbindungen eines Parlamentariers ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and PersonNumber eq <personNumber>", 'Personenfilter'),
    ('$select', 'PersonNumber,InterestName,InterestTypeText,...', 'Feldauswahl'),
    ('$format', 'json', 'Antwortformat'),
])

# --- 5.12 PersonOccupation ---
doc.add_heading('5.12 PersonOccupation (Berufe / Tätigkeiten)', level=2)
add_endpoint_box(doc, 'GET', '/PersonOccupation',
                 'Ruft die beruflichen Tätigkeiten eines Parlamentariers ab.')
add_query_params_table(doc, [
    ('$filter', "Language eq 'DE' and PersonNumber eq <personNumber>", 'Personenfilter'),
    ('$select', 'PersonNumber,OccupationName,Employer,JobTitle', 'Feldauswahl'),
    ('$format', 'json', 'Antwortformat'),
])

# ============================================================
# 6. DATENMODELLE
# ============================================================
doc.add_page_break()
doc.add_heading('6. Datenmodelle (DTOs)', level=1)
doc.add_paragraph(
    'Die folgenden Tabellen beschreiben alle verfügbaren Attribute der API-Antworten. '
    'Alle Felder ausser den mit "Ja" markierten Pflichtfeldern sind optional (nullable).'
)

# SessionDTO
doc.add_heading('6.1 SessionDTO', level=2)
add_field_table(doc, [
    ('ID', 'Int', 'Ja', 'Eindeutige Session-ID'),
    ('SessionNumber', 'Int', 'Nein', 'Sessionsnummer'),
    ('SessionName', 'String', 'Nein', 'Name der Session (z.B. "Wintersession 2024")'),
    ('Abbreviation', 'String', 'Nein', 'Abkürzung der Session'),
    ('StartDate', 'DateTime', 'Nein', 'Startdatum der Session (/Date(ms)/ Format)'),
    ('EndDate', 'DateTime', 'Nein', 'Enddatum der Session'),
    ('Title', 'String', 'Nein', 'Titel der Session'),
    ('Type', 'Int', 'Nein', 'Sessionstyp (numerisch, gemappt als SessionType)'),
    ('TypeName', 'String', 'Nein', 'Beschreibung des Sessionstyps'),
    ('LegislativePeriodNumber', 'Int', 'Nein', 'Nummer der Legislaturperiode'),
])

# GeschaeftDTO
doc.add_heading('6.2 GeschaeftDTO (Business)', level=2)
add_field_table(doc, [
    ('ID', 'Int', 'Ja', 'Eindeutige Geschäfts-ID'),
    ('BusinessShortNumber', 'String', 'Nein', 'Kurzbezeichnung (z.B. "24.3456")'),
    ('Title', 'String', 'Nein', 'Titel des Geschäfts'),
    ('BusinessTypeName', 'String', 'Nein', 'Geschäftstyp (z.B. "Motion", "Interpellation")'),
    ('BusinessTypeAbbreviation', 'String', 'Nein', 'Abkürzung des Geschäftstyps (z.B. "Mo.", "Ip.")'),
    ('BusinessStatusText', 'String', 'Nein', 'Aktueller Status des Geschäfts'),
    ('BusinessStatusDate', 'DateTime', 'Nein', 'Datum des letzten Statuswechsels'),
    ('SubmissionDate', 'DateTime', 'Nein', 'Einreichungsdatum'),
    ('SubmittedBy', 'String', 'Nein', 'Eingereicht von (Ratsmitglied oder Fraktion)'),
    ('Description', 'String', 'Nein', 'Ausführliche Beschreibung / Text des Vorstosses'),
    ('SubmissionCouncilName', 'String', 'Nein', 'Einreichungsrat (Nationalrat/Ständerat)'),
    ('ResponsibleDepartmentName', 'String', 'Nein', 'Zuständiges Departement'),
    ('ResponsibleDepartmentAbbreviation', 'String', 'Nein', 'Abkürzung des Departements (z.B. "EFD")'),
    ('TagNames', 'String', 'Nein', 'Schlagwörter / Themenbereiche'),
    ('Modified', 'DateTime', 'Nein', 'Letzte Änderung (für inkrementelle Synchronisation)'),
])

# TranscriptDTO
doc.add_heading('6.3 TranscriptDTO (Wortmeldung)', level=2)
add_field_table(doc, [
    ('ID', 'String', 'Nein', 'Eindeutige Transkript-ID'),
    ('PersonNumber', 'Int', 'Nein', 'Personennummer des Redners'),
    ('SpeakerFullName', 'String', 'Nein', 'Vollständiger Name des Redners'),
    ('SpeakerFunction', 'String', 'Nein', 'Funktion (z.B. "Bundesrat", "Berichterstatter")'),
    ('Text', 'String', 'Nein', 'Redetext (HTML-formatiert)'),
    ('MeetingDate', 'String', 'Nein', 'Datum der Sitzung'),
    ('Start', 'DateTime', 'Nein', 'Startzeitpunkt der Wortmeldung'),
    ('End', 'DateTime', 'Nein', 'Endzeitpunkt der Wortmeldung'),
    ('CouncilName', 'String', 'Nein', 'Ratsbezeichnung'),
    ('ParlGroupAbbreviation', 'String', 'Nein', 'Fraktionsabkürzung (z.B. "S", "V", "M-E")'),
    ('CantonAbbreviation', 'String', 'Nein', 'Kantonsabkürzung (z.B. "ZH", "BE")'),
    ('SortOrder', 'Int', 'Nein', 'Reihenfolge innerhalb der Debatte'),
    ('Type', 'Int', 'Nein', 'Typ der Wortmeldung (1 = Rede/Speech)'),
])

# SubjectBusinessDTO
doc.add_heading('6.4 SubjectBusinessDTO', level=2)
add_field_table(doc, [
    ('IdSubject', 'String', 'Nein', 'ID des Tagesordnungspunkts'),
    ('BusinessNumber', 'Int', 'Nein', 'Geschäftsnummer'),
    ('BusinessShortNumber', 'String', 'Nein', 'Kurzbezeichnung des Geschäfts'),
    ('Title', 'String', 'Nein', 'Titel des Geschäfts'),
    ('SortOrder', 'Int', 'Nein', 'Sortierreihenfolge'),
])

# MeetingDTO
doc.add_heading('6.5 MeetingDTO (Sitzung)', level=2)
add_field_table(doc, [
    ('ID', 'String', 'Nein', 'Eindeutige Sitzungs-ID'),
    ('MeetingNumber', 'Int', 'Nein', 'Sitzungsnummer'),
    ('IdSession', 'Int', 'Nein', 'ID der zugehörigen Session'),
    ('Council', 'Int', 'Nein', 'Ratsnummer (numerisch)'),
    ('CouncilName', 'String', 'Nein', 'Ratsbezeichnung (Nationalrat/Ständerat)'),
    ('CouncilAbbreviation', 'String', 'Nein', 'Ratsabkürzung (NR/SR)'),
    ('Date', 'DateTime', 'Nein', 'Datum der Sitzung'),
    ('Begin', 'String', 'Nein', 'Beginn der Sitzung (Uhrzeit)'),
    ('MeetingOrderText', 'String', 'Nein', 'Tagesordnung / Sitzungsbezeichnung'),
    ('SortOrder', 'Int', 'Nein', 'Sortierreihenfolge'),
    ('SessionName', 'String', 'Nein', 'Name der zugehörigen Session'),
])

# SubjectDTO
doc.add_heading('6.6 SubjectDTO (Tagesordnungspunkt)', level=2)
add_field_table(doc, [
    ('ID', 'String', 'Nein', 'Eindeutige Subject-ID'),
    ('IdMeeting', 'String', 'Nein', 'ID der zugehörigen Sitzung'),
    ('SortOrder', 'Int', 'Nein', 'Reihenfolge im Tagesordnungspunkt'),
    ('VerbalixOid', 'Int', 'Nein', 'Verbalix-Objekt-ID (für Transkriptverknüpfung)'),
])

# BusinessRoleDTO
doc.add_heading('6.7 BusinessRoleDTO (Geschäftsrolle)', level=2)
add_field_table(doc, [
    ('ID', 'String', 'Nein', 'Eindeutige Rollen-ID'),
    ('Role', 'Int', 'Nein', 'Rollentyp (7 = Urheber/Originator)'),
    ('RoleName', 'String', 'Nein', 'Bezeichnung der Rolle'),
    ('BusinessNumber', 'Int', 'Nein', 'Zugehörige Geschäftsnummer'),
    ('MemberCouncilNumber', 'Int', 'Nein', 'Personennummer des Ratsmitglieds'),
])

# ParlamentarierDTO
doc.add_heading('6.8 ParlamentarierDTO (MemberCouncil)', level=2)
doc.add_paragraph('Basisfelder (immer verfügbar):', style='List Bullet')
add_field_table(doc, [
    ('ID', 'Int', 'Nein', 'Eindeutige ID'),
    ('PersonNumber', 'Int', 'Nein', 'Personennummer'),
    ('FirstName', 'String', 'Nein', 'Vorname'),
    ('LastName', 'String', 'Nein', 'Nachname'),
    ('PartyAbbreviation', 'String', 'Nein', 'Parteiabkürzung (z.B. "SP", "SVP", "FDP")'),
    ('ParlGroupAbbreviation', 'String', 'Nein', 'Fraktionsabkürzung'),
    ('CantonAbbreviation', 'String', 'Nein', 'Kantonsabkürzung'),
    ('CouncilName', 'String', 'Nein', 'Ratsbezeichnung'),
    ('CouncilAbbreviation', 'String', 'Nein', 'Ratsabkürzung (NR/SR)'),
    ('Active', 'Bool', 'Nein', 'Aktives Ratsmitglied (true/false)'),
], caption='Basisfelder')

add_field_table(doc, [
    ('DateOfBirth', 'DateTime', 'Nein', 'Geburtsdatum'),
    ('MaritalStatusText', 'String', 'Nein', 'Zivilstand'),
    ('NumberOfChildren', 'Int', 'Nein', 'Anzahl Kinder'),
    ('BirthPlace_City', 'String', 'Nein', 'Geburtsort (Stadt)'),
    ('BirthPlace_Canton', 'String', 'Nein', 'Geburtsort (Kanton)'),
    ('Citizenship', 'String', 'Nein', 'Bürgerort'),
    ('DateJoining', 'DateTime', 'Nein', 'Eintrittsdatum in den Rat'),
    ('DateLeaving', 'DateTime', 'Nein', 'Austrittsdatum aus dem Rat'),
    ('DateElection', 'DateTime', 'Nein', 'Wahldatum'),
    ('MilitaryRankText', 'String', 'Nein', 'Militärischer Grad'),
    ('PartyName', 'String', 'Nein', 'Vollständiger Parteiname'),
    ('ParlGroupName', 'String', 'Nein', 'Vollständiger Fraktionsname'),
    ('CantonName', 'String', 'Nein', 'Vollständiger Kantonsname'),
    ('Nationality', 'String', 'Nein', 'Nationalität'),
], caption='Detailfelder (zusätzlich via $select)')

# VoteDTO
doc.add_heading('6.9 VoteDTO (Abstimmung)', level=2)
add_field_table(doc, [
    ('ID', 'Int', 'Ja', 'Eindeutige Abstimmungs-ID'),
    ('BusinessNumber', 'Int', 'Nein', 'Zugehörige Geschäftsnummer'),
    ('BusinessShortNumber', 'String', 'Nein', 'Kurzbezeichnung des Geschäfts'),
    ('BillTitle', 'String', 'Nein', 'Titel der Vorlage'),
    ('IdSession', 'Int', 'Nein', 'ID der zugehörigen Session'),
    ('Subject', 'String', 'Nein', 'Abstimmungsgegenstand'),
    ('MeaningYes', 'String', 'Nein', 'Bedeutung einer Ja-Stimme'),
    ('MeaningNo', 'String', 'Nein', 'Bedeutung einer Nein-Stimme'),
    ('VoteEnd', 'DateTime', 'Nein', 'Zeitpunkt des Abstimmungsendes'),
    ('Modified', 'DateTime', 'Nein', 'Letzte Änderung'),
])

# VotingDTO
doc.add_heading('6.10 VotingDTO (Stimmabgabe)', level=2)
add_field_table(doc, [
    ('ID', 'Int', 'Ja', 'Eindeutige ID der Stimmabgabe'),
    ('IdVote', 'Int', 'Nein', 'ID der zugehörigen Abstimmung'),
    ('PersonNumber', 'Int', 'Nein', 'Personennummer des Ratsmitglieds'),
    ('Decision', 'Int', 'Nein', 'Entscheid (1=Ja, 2=Nein, 3=Enthaltung, 4=Nicht teilgenommen, 5=Entschuldigt, 6=Präsident)'),
    ('DecisionText', 'String', 'Nein', 'Textuelle Beschreibung des Entscheids'),
])

# PersonInterestDTO
doc.add_heading('6.11 PersonInterestDTO (Interessensbindung)', level=2)
add_field_table(doc, [
    ('PersonNumber', 'Int', 'Nein', 'Personennummer'),
    ('InterestName', 'String', 'Nein', 'Name der Organisation / Interessensbindung'),
    ('InterestTypeText', 'String', 'Nein', 'Art der Interessensbindung'),
    ('FunctionInAgencyText', 'String', 'Nein', 'Funktion in der Organisation'),
    ('Paid', 'Bool', 'Nein', 'Bezahlt (true/false)'),
    ('OrganizationTypeText', 'String', 'Nein', 'Art der Organisation'),
])

# PersonOccupationDTO
doc.add_heading('6.12 PersonOccupationDTO (Beruf)', level=2)
add_field_table(doc, [
    ('PersonNumber', 'Int', 'Nein', 'Personennummer'),
    ('OccupationName', 'String', 'Nein', 'Berufsbezeichnung'),
    ('Employer', 'String', 'Nein', 'Arbeitgeber'),
    ('JobTitle', 'String', 'Nein', 'Stellenbezeichnung'),
])

# ============================================================
# 7. OData Response-Wrapper
# ============================================================
doc.add_page_break()
doc.add_heading('7. OData Response-Wrapper', level=1)

doc.add_heading('7.1 Listen-Response (ODataResponse<T>)', level=2)
doc.add_paragraph('Antwortstruktur für Abfragen, die mehrere Ergebnisse liefern:')

p = doc.add_paragraph()
run = p.add_run('''{
  "d": {
    "results": [
      { ...item1... },
      { ...item2... }
    ],
    "__next": "https://ws.parlament.ch/odata.svc/...?$skiptoken=..."
  }
}''')
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_field_table(doc, [
    ('d.results', 'Array<T>', 'Ja', 'Liste der Ergebnisobjekte'),
    ('d.__next', 'String', 'Nein', 'URL zur nächsten Seite (Paginierung)'),
])

doc.add_paragraph('Alternative Struktur (einfaches Array):')
p = doc.add_paragraph()
run = p.add_run('''{
  "d": [
    { ...item1... },
    { ...item2... }
  ]
}''')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('7.2 Einzel-Response (ODataSingleResponse<T>)', level=2)
doc.add_paragraph('Antwortstruktur für Abfragen, die ein einzelnes Ergebnis liefern:')

p = doc.add_paragraph()
run = p.add_run('''{
  "d": { ...item... }
}''')
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_field_table(doc, [
    ('d', 'T', 'Ja', 'Das einzelne Ergebnisobjekt'),
])

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/thomassussli/Documents/XCode/Politik/Swiss_Parliament_OData_API_Spezifikation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
